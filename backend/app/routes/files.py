from io import BytesIO
from pathlib import Path
from typing import List

from docx import Document
from fastapi import APIRouter, File, HTTPException, UploadFile
from pypdf import PdfReader


router = APIRouter()

MAX_FILE_SIZE = 10 * 1024 * 1024

ALLOWED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
}


def extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))

    pages = []

    for page in reader.pages:
        text = page.extract_text() or ""

        if text.strip():
            pages.append(text.strip())

    return "\n\n".join(pages)


def extract_docx_text(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))

    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = []

            for cell in row.cells:
                text = cell.text.strip()

                if text:
                    row_text.append(text)

            if row_text:
                parts.append(" | ".join(row_text))

    return "\n".join(parts)


def extract_text_file(file_bytes: bytes) -> str:
    return file_bytes.decode(
        "utf-8",
        errors="replace"
    )


def extract_file_text(
    filename: str,
    file_bytes: bytes
) -> str:

    extension = Path(filename).suffix.lower()

    if extension == ".pdf":
        return extract_pdf_text(file_bytes)

    if extension == ".docx":
        return extract_docx_text(file_bytes)

    if extension in {".txt", ".md"}:
        return extract_text_file(file_bytes)

    raise ValueError(
        f"Unsupported file type: {extension}"
    )


@router.post("/api/files/upload")
async def upload_files(
    files: List[UploadFile] = File(...)
):
    if not files:
        raise HTTPException(
            status_code=400,
            detail="No files were uploaded."
        )

    results = []

    for uploaded_file in files:

        filename = uploaded_file.filename or "unnamed"

        extension = Path(filename).suffix.lower()

        if extension not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"{filename} is not supported. "
                    "Supported files: PDF, DOCX, TXT, MD."
                )
            )

        try:
            file_bytes = await uploaded_file.read()

            if len(file_bytes) == 0:
                raise HTTPException(
                    status_code=400,
                    detail=f"{filename} is empty."
                )

            if len(file_bytes) > MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=413,
                    detail=(
                        f"{filename} is too large. "
                        "Maximum file size is 10 MB."
                    )
                )

            text = extract_file_text(
                filename,
                file_bytes
            )

            results.append({
                "name": filename,
                "content_type": uploaded_file.content_type,
                "size": len(file_bytes),
                "characters": len(text),
                "text": text,
            })

        except HTTPException:
            raise

        except Exception as error:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"Could not read {filename}: "
                    f"{str(error)}"
                )
            )

        finally:
            await uploaded_file.close()

    combined_text_parts = []

    for file_data in results:
        combined_text_parts.append(
            f"===== FILE: {file_data['name']} =====\n"
            f"{file_data['text']}"
        )

    combined_text = "\n\n".join(
        combined_text_parts
    )

    return {
        "success": True,
        "count": len(results),
        "files": results,
        "combined_text": combined_text,
    }